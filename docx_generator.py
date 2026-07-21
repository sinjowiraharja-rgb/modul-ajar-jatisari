from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def add_heading(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 128)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(str(text))
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(str(item))
        run.font.name = 'Arial'
        run.font.size = Pt(11)


def generate_docx(modul, output_path):
    """
    Generate dokumen Word (.docx) dari data modul ajar.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Judul
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"MODUL AJAR\n{modul['mapel'].upper()}")
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{modul['topik']}\nKelas {modul['kelas']} | Fase {modul['fase']}")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    subtitle.paragraph_format.space_after = Pt(20)

    # A. INFORMASI UMUM
    add_heading(doc, 'A. INFORMASI UMUM', level=1)

    tabel_identitas = doc.add_table(rows=1, cols=2)
    tabel_identitas.style = 'Table Grid'
    hdr_cells = tabel_identitas.rows[0].cells
    hdr_cells[0].text = 'Komponen'
    hdr_cells[1].text = 'Keterangan'

    identitas = [
        ('Nama Penyusun', modul.get('nama_guru', '')),
        ('Kepala Sekolah', modul.get('nama_kepala_sekolah', '')),
        ('Satuan Pendidikan', modul.get('nama_sekolah', '')),
        ('Fase / Kelas', f"{modul.get('fase', '')} / {modul.get('kelas', '')}"),
        ('Mata Pelajaran', modul.get('mapel', '')),
        ('Topik / Materi', modul.get('topik', '')),
        ('Alokasi Waktu', modul.get('alokasi_waktu', '')),
        ('Minggu Pelaksanaan', modul.get('minggu_pelaksanaan', '')),
        ('Model Pembelajaran', modul.get('model_pembelajaran', '')),
        ('Pendekatan', 'Deep Learning' if modul.get('pendekatan') == 'deep_learning' else 'S.T.E.A.M'),
    ]

    for label, value in identitas:
        row_cells = tabel_identitas.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)

    doc.add_paragraph()

    add_heading(doc, 'Kompetensi Awal', level=2)
    add_paragraph(doc, modul.get('kompetensi_awal', '-'))

    add_heading(doc, 'Profil Pelajar Pancasila', level=2)
    dimensi = modul.get('dimensi_pancasila', '')
    if isinstance(dimensi, str) and dimensi:
        dimensi_list = [d.strip() for d in dimensi.split(',') if d.strip()]
        add_bullet_list(doc, dimensi_list)
    elif isinstance(dimensi, list):
        add_bullet_list(doc, dimensi)
    else:
        add_paragraph(doc, '-')

    # B. KOMPONEN INTI
    add_heading(doc, 'B. KOMPONEN INTI', level=1)

    add_heading(doc, 'Tujuan Pembelajaran', level=2)
    add_paragraph(doc, modul.get('tujuan_pembelajaran', '-'))

    add_heading(doc, 'Pemahaman Bermakna', level=2)
    add_paragraph(doc, modul.get('pemahaman_bermakna', '-'))

    add_heading(doc, 'Pertanyaan Pemantik', level=2)
    add_paragraph(doc, modul.get('pertanyaan_pemantik', '-'))

    add_heading(doc, 'Persiapan Pembelajaran', level=2)
    add_paragraph(doc, modul.get('persiapan_pembelajaran', '-'))

    add_heading(doc, 'Sarana dan Prasarana', level=2)
    add_paragraph(doc, modul.get('sarana_prasarana', '-'))

    add_heading(doc, 'Target Peserta Didik', level=2)
    add_paragraph(doc, modul.get('target_peserta_didik', '-'))

    add_heading(doc, 'Langkah-Langkah Pembelajaran', level=2)
    add_paragraph(doc, modul.get('langkah_pembelajaran', '-'))

    add_heading(doc, 'Asesmen', level=2)
    add_paragraph(doc, modul.get('asesmen', '-'))

    add_heading(doc, 'Pengayaan', level=2)
    add_paragraph(doc, modul.get('pengayaan', '-'))

    add_heading(doc, 'Remedial', level=2)
    add_paragraph(doc, modul.get('remedial', '-'))

    # C. LAMPIRAN
    add_heading(doc, 'C. LAMPIRAN', level=1)

    add_heading(doc, 'Lembar Kerja Peserta Didik (LKPD)', level=2)
    add_paragraph(doc, modul.get('lkpd', '-'))

    add_heading(doc, 'Bahan Bacaan Guru dan Murid', level=2)
    add_paragraph(doc, modul.get('bahan_bacaan', '-'))

    add_heading(doc, 'Glosarium', level=2)
    add_paragraph(doc, modul.get('glosarium', '-'))

    add_heading(doc, 'Daftar Pustaka', level=2)
    add_paragraph(doc, modul.get('daftar_pustaka', '-'))

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Modul Ajar Jatisari | Dibuat otomatis melalui aplikasi Modul Ajar Jatisari")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(output_path)
    return output_path
